import sys
sys.stdout.reconfigure(encoding="utf-8")
import docx

doc = docx.Document(r"链接.docx")

print("=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        print(f"P{i}: {t}")

print("\n=== TABLES ===")
for ti, table in enumerate(doc.tables):
    print(f"\n--- Table {ti} ---")
    for ri, row in enumerate(table.rows):
        cells = [c.text.strip() for c in row.cells]
        print(f"  Row {ri}: {' | '.join(cells)}")

# Also extract hyperlinks
from docx.opc.constants import RELATIONSHIP_TYPE as RT
print("\n=== HYPERLINKS ===")
for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        if run._r.xml.count("w:hyperlink") > 0:
            import re
            hrefs = re.findall(r'r:id="([^"]+)"', run._r.xml)
            for rid in hrefs:
                try:
                    rel = doc.part.rels[rid]
                    print(f"P{i}: {rel.target_ref}")
                except:
                    pass

# Check all relationships
print("\n=== ALL RELATIONSHIPS ===")
for rid, rel in doc.part.rels.items():
    if "hyperlink" in str(rel.reltype):
        print(f"{rid}: {rel.target_ref}")
