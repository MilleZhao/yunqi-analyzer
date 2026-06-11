import sys, json, os, re
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

workdir = r"WORKDIR"

# Load both datasets
with open(os.path.join(workdir, "links_parsed.json"), "r", encoding="utf-8") as f:
    links_data = json.load(f)

with open(os.path.join(workdir, "douyin_results.json"), "r", encoding="utf-8") as f:
    dy_raw = json.load(f)

# --- Helper: add clickable hyperlink to paragraph ---
def add_link(paragraph, url, text=None):
    if not text:
        text = url
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    for tag, val in [("w:color", "0563C1"), ("w:u", "single"), ("w:sz", "20"), ("w:rFonts", "Calibri")]:
        el = OxmlElement(tag)
        if tag == "w:rFonts":
            el.set(qn("w:ascii"), val)
        else:
            el.set(qn("w:val"), val)
        rPr.append(el)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph

# --- Build Douyin accounts from scraped data ---
children_accounts = []
general_accounts = []
seen_uids = set()

for group in dy_raw.get("videos", []):
    for api in group.get("apiResponses", []):
        d = api.get("data", {})
        for bd in d.get("business_data", []):
            info = None
            bdata = bd.get("data", {})
            if bdata.get("aweme_info"):
                info = bdata["aweme_info"]
            elif bdata.get("user_info"):
                info = bdata["user_info"]
            if info:
                author = info.get("author", info)
                uid = author.get("uid", "")
                if uid and uid not in seen_uids:
                    seen_uids.add(uid)
                    acc = {
                        "nickname": author.get("nickname", "") or author.get("unique_id", ""),
                        "uid": uid,
                        "signature": (author.get("signature", "") or "").strip().replace("\n", " | "),
                        "link": f"https://www.douyin.com/user/{author.get('sec_uid', uid)}"
                    }
                    # Categorize
                    combined = acc["nickname"] + " " + acc["signature"]
                    children_kw = ["童书","少儿","儿童","绘本","小朋友","宝宝","育儿","早教","幼儿","亲子","妈妈","麻麻","爸爸","养娃","伴读","园长","小学生","启蒙","教小学"]
                    shop_kw = ["专营店","旗舰店","书店","图书店","童学"]
                    key_names = ["青葫芦","小爱童年","老夏挑绘本","安东尼","爱读书的硕硕","米团妈妈","叶子老师","安安小朋友","甜甜园长","开心鱼头","嘟嘟妈","恺欣麻麻","海宝兄弟","海淀大牛","深圳童学","烁妈伴读","三小辫儿园长","悦悦麻麻","大琪是来宝妈","鹏妈陪读","三宝说英语","央视网阅读","喜哥图书店","维乾图书","小嘉啊"]
                    
                    is_children = any(kw in combined for kw in children_kw) or any(kw in acc["nickname"] for kw in key_names) or any(kw in combined for kw in shop_kw)
                    if is_children:
                        children_accounts.append(acc)
                    else:
                        general_accounts.append(acc)

print(f"Children accounts: {len(children_accounts)}")
print(f"General accounts: {len(general_accounts)}")

# --- CREATE DOCUMENT ---
doc = Document()

# Page setup
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Styles
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.2

for lvl, (sz, clr) in {1: (16, "2E74B5"), 2: (13, "2E74B5"), 3: (12, "1F4D78")}.items():
    hs = doc.styles[f"Heading {lvl}"]
    hs.font.name = "Calibri"
    hs.font.size = Pt(sz)
    hs.font.color.rgb = RGBColor.from_string(clr)
    hs.font.bold = True

# === TITLE ===
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run("少儿图书 / 图书 相关账号与链接汇总")
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = RGBColor.from_string("0B2545")
r.font.name = "Calibri"

sub = doc.add_paragraph()
r = sub.add_run("数据来源：小红书 + 抖音  |  2026年6月9日")
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
sub.paragraph_format.space_after = Pt(16)

# === Helper functions ===
def add_account_item(doc, nickname, link, signature=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(nickname)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor.from_string("0B2545")
    if link:
        p.add_run("  ").font.size = Pt(9)
        add_link(p, link, "打开链接")
    if signature:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(6)
        p2.paragraph_format.left_indent = Inches(0.25)
        r2 = p2.add_run(signature[:200])
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        r2.font.name = "Calibri"

def add_section_header(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_link_item(doc, label, url, desc=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if label:
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"
        p.add_run("  ").font.size = Pt(8)
    add_link(p, url, url)
    if desc:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(6)
        p2.paragraph_format.left_indent = Inches(0.25)
        r2 = p2.add_run(desc[:200])
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        r2.font.name = "Calibri"

# ============ PART 1: XIAOHONGSHU ============
add_section_header(doc, "小红书", level=1)

# 1.1 Children's books
add_section_header(doc, "少儿童书", level=2)
xhs_c = links_data["xhs_children"]
print(f"XHS children: {len(xhs_c)}")
for item in xhs_c:
    label = item["account"] if item["account"] else "小红书链接"
    add_link_item(doc, f"【{label}】", item["url"], item.get("desc", ""))

# 1.2 General books
add_section_header(doc, "图书", level=2)
xhs_b = links_data["xhs_books"]
print(f"XHS books: {len(xhs_b)}")
for item in xhs_b:
    label = item["account"] if item["account"] else "小红书链接"
    add_link_item(doc, f"【{label}】", item["url"], item.get("desc", ""))

# ============ PART 2: DOUYIN ACCOUNTS ============
add_section_header(doc, "抖音 — 搜到的账号（带主页链接）", level=1)

add_section_header(doc, "少儿童书 / 童书 / 绘本 相关账号", level=2)
for acc in children_accounts:
    add_account_item(doc, acc["nickname"], acc["link"], acc.get("signature", ""))

add_section_header(doc, "图书 / 阅读 相关账号", level=2)
for acc in general_accounts:
    add_account_item(doc, acc["nickname"], acc["link"], acc.get("signature", ""))

# ============ PART 3: DOUYIN VIDEO LINKS (from 链接.docx) ============
add_section_header(doc, "抖音 — 链接.docx 中的视频链接", level=1)

dy_c = links_data["dy_children"]
if dy_c:
    add_section_header(doc, "少儿童书 视频", level=2)
    for item in dy_c:
        add_link_item(doc, "", item["url"], item.get("desc", ""))

dy_b = links_data["dy_books"]
if dy_b:
    add_section_header(doc, "图书 视频", level=2)
    for item in dy_b:
        add_link_item(doc, "", item["url"], item.get("desc", ""))

# === Footer ===
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("少儿图书 / 图书 账号与链接汇总  |  2026.06.09")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# === SAVE ===
output = os.path.join(workdir, "少儿图书_图书_账号链接汇总.docx")
doc.save(output)
print(f"\nSaved: {output}")
print(f"Total: XHS {len(xhs_c)+len(xhs_b)} links + DY accounts {len(children_accounts)+len(general_accounts)} + DY videos {len(dy_c)+len(dy_b)}")
