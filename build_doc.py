import json
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Load data
workdir = r"C:\Users\LENOVO\Documents\云阙智能1"
with open(os.path.join(workdir, "douyin_results.json"), "r", encoding="utf-8") as f:
    raw = json.load(f)

# Collect all unique accounts from API data
accounts = {}
all_videos = []

for group in raw.get("videos", []):
    keyword = group["keyword"]
    for api in group.get("apiResponses", []):
        d = api.get("data", {})
        for bd in d.get("business_data", []):
            info = None
            bdata = bd.get("data", {})
            if bdata.get("aweme_info"):
                info = bdata["aweme_info"]
            elif bdata.get("user_info"):
                info = bdata["user_info"]
            if info:
                author = info.get("author", info)
                uid = author.get("uid", "")
                if uid and uid not in accounts:
                    accounts[uid] = {
                        "nickname": author.get("nickname", "") or author.get("unique_id", ""),
                        "uid": uid,
                        "signature": (author.get("signature", "") or "").strip().replace("\n", " | "),
                        "link": f"https://www.douyin.com/user/{author.get('sec_uid', uid)}"
                    }
                if bdata.get("aweme_info"):
                    all_videos.append({
                        "keyword": keyword,
                        "desc": bdata["aweme_info"].get("desc", ""),
                        "author": author.get("nickname", ""),
                        "video_link": f"https://www.douyin.com/video/{bdata['aweme_info']['aweme_id']}",
                        "author_link": f"https://www.douyin.com/user/{author.get('sec_uid', author.get('uid', ''))}"
                    })

# Categorize accounts
children_accounts = []  # Kids/children book related
general_accounts = []   # General books / reading

children_keywords = ["童书", "少儿", "儿童", "绘本", "小朋友", "宝宝", "育儿", "早教", "幼", "童", "亲子", "妈妈", "麻麻", "爸爸", "养娃", "崽", "萌", "伴读", "园长", "小学生"]
shop_keywords = ["图书", "书", "阅读", "read", "Book", "book"]

for uid, acc in accounts.items():
    nick = acc["nickname"]
    sig = acc["signature"]
    combined = nick + " " + sig
    
    is_children = any(kw in combined for kw in children_keywords)
    is_shop = any(kw in combined for kw in ["专营店", "旗舰店", "书店", "图书店"] + ["图书"] if not acc["nickname"] in ["维乾图书", "喜哥图书店"])
    
    if is_children or is_shop or any(kw in acc["nickname"] for kw in ["青葫芦", "小爱童年", "老夏挑绘本", "安东尼", "爱读书的硕硕", "米团妈妈", "叶子老师", "安安小朋友", "甜甜园长", "开心鱼头", "嘟嘟妈", "恺欣麻麻", "海宝兄弟", "海淀大牛", "深圳童学", "烁妈伴读", "三小辫儿园长", "悦悦麻麻", "大琪是来宝妈", "鹏妈陪读", "三宝说英语", "央视网阅读"]):
        children_accounts.append(acc)
    else:
        general_accounts.append(acc)

# Also add page-scraped author names
page_authors = set()
for group in raw.get("videos", []):
    for r in group.get("results", []):
        name = (r.get("author") or r.get("name") or "").strip()
        # Clean date suffixes
        import re
        name = re.sub(r'\d{4}\.\d{1,2}\.\d{1,2}', '', name).strip()
        name = re.sub(r'^[昨天前天今天]+', '', name).strip()
        if name and len(name) > 1 and name not in ["未知", "robin"]:
            page_authors.add(name)

# Filter out names already in accounts
existing_names = {a["nickname"] for a in accounts.values()}
# Also match by partial: remove emoji and compare
extra_authors = []
for name in sorted(page_authors):
    # Check if this name or a cleaned version matches existing
    clean = name.replace("🧸", "").replace("🧡", "").replace("💛", "").strip()
    if not any(clean in ex or ex in clean for ex in existing_names):
        extra_authors.append(name)

# === CREATE DOCUMENT ===
doc = Document()

# Page setup
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Style setup
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.25

# Heading styles
for level, (size, color) in {1: (16, "2E74B5"), 2: (13, "2E74B5"), 3: (12, "1F4D78")}.items():
    hstyle = doc.styles[f"Heading {level}"]
    hstyle.font.name = "Calibri"
    hstyle.font.size = Pt(size)
    hstyle.font.color.rgb = RGBColor.from_string(color)
    hstyle.font.bold = True

# === TITLE ===
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title.add_run("抖音少儿图书 / 图书 相关账号汇总")
run.font.size = Pt(20)
run.font.bold = True
run.font.color.rgb = RGBColor.from_string("0B2545")
run.font.name = "Calibri"

# Subtitle / date
sub = doc.add_paragraph()
run = sub.add_run("搜索日期：2026年6月9日  |  数据来源：抖音搜索（少儿图书 / 童书 / 绘本推荐 / 图书）")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = "Calibri"
sub.paragraph_format.space_after = Pt(12)

# === Helper: add hyperlink paragraph ===
def add_account_para(doc, nickname, link, signature=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    
    # Account name as bold
    run_name = p.add_run(nickname)
    run_name.bold = True
    run_name.font.size = Pt(10)
    run_name.font.name = "Calibri"
    run_name.font.color.rgb = RGBColor.from_string("0B2545")
    
    # Link
    if link:
        p.add_run("  ").font.size = Pt(9)
        add_hyperlink(p, link, "打开主页", "0563C1")
    
    # Signature
    if signature:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(6)
        p2.paragraph_format.left_indent = Inches(0.25)
        run_sig = p2.add_run(signature[:200])
        run_sig.font.size = Pt(8.5)
        run_sig.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run_sig.font.name = "Calibri"
    return p

def add_hyperlink(paragraph, url, text, color="0563C1"):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    rPr.append(sz)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rPr.append(rFonts)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph

# === SECTION 1: Children's Book Accounts ===
doc.add_heading("一、少儿图书 / 童书 / 绘本 相关账号", level=1)

# Sort by relevance
children_sorted = sorted(children_accounts, key=lambda x: (
    0 if "青葫芦" in x["nickname"] else
    1 if "旗舰店" in x["nickname"] or "专营店" in x["nickname"] else
    2 if "老夏" in x["nickname"] or "安东尼" in x["nickname"] else
    3, x["nickname"]
))

doc.add_paragraph(f"共 {len(children_sorted)} 个账号", style="Normal").runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

for acc in children_sorted:
    add_account_para(doc, acc["nickname"], acc["link"], acc["signature"])

# === SECTION 2: General Book / Reading Accounts ===
doc.add_heading("二、图书 / 阅读 相关账号", level=1)
general_sorted = sorted(general_accounts, key=lambda x: x["nickname"])
doc.add_paragraph(f"共 {len(general_sorted)} 个账号", style="Normal").runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

for acc in general_sorted:
    add_account_para(doc, acc["nickname"], acc["link"], acc["signature"])

# === SECTION 3: Additional Page-Scraped Authors ===
if extra_authors:
    doc.add_heading("三、页面抓取到的更多账号名（无完整信息）", level=1)
    doc.add_paragraph(f"共 {len(extra_authors)} 个，来自搜索结果页直接提取", style="Normal").runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    # Split into columns using a simple format
    p = doc.add_paragraph()
    for i, name in enumerate(extra_authors):
        if i > 0:
            p.add_run("  |  ").font.size = Pt(8)
        run = p.add_run(name)
        run.font.size = Pt(9)
        run.font.name = "Calibri"

# === SECTION 4: Sample Videos ===
doc.add_heading("四、部分视频链接（按关键词）", level=1)

video_by_keyword = {}
for v in all_videos:
    kw = v["keyword"]
    if kw not in video_by_keyword:
        video_by_keyword[kw] = []
    if len(video_by_keyword[kw]) < 5:
        video_by_keyword[kw].append(v)

for kw, vids in video_by_keyword.items():
    doc.add_heading(kw, level=3)
    for v in vids[:3]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        desc = v["desc"][:100] if v["desc"] else "(无描述)"
        run = p.add_run(f"{v['author']}: {desc}")
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(4)
        add_hyperlink(p2, v["video_link"], v["video_link"], "0563C1")
        # size set in hyperlink

# Footer
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("抖音账号汇总  |  2026.06.09")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# Save
output_path = os.path.join(workdir, "抖音图书账号汇总.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
print(f"Accounts: {len(children_sorted)} children + {len(general_sorted)} general = {len(children_sorted) + len(general_sorted)} total")
print(f"Extra authors: {len(extra_authors)}")
print(f"Videos: {len(all_videos)}")
