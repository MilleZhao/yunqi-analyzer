import sys, json, re, os
sys.stdout.reconfigure(encoding="utf-8")
import docx as dx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

workdir = r"WORKDIR"
links_path = r"C:\Users\LENOVO\Documents\CodexProjects\Project018-云阙智能1\链接.docx"

# === READ 链接.docx ===
links_doc = dx.Document(links_path)

xhs_children = []  # 小红书少儿童书
xhs_books = []     # 小红书图书
dy_children = []   # 抖音少儿童书
dy_books = []      # 抖音图书

current_section = None
for p in links_doc.paragraphs:
    t = p.text.strip()
    if not t:
        continue
    
    if "少儿童书——小红书" in t:
        current_section = "xhs_children"
        continue
    elif "图书——小红书" in t:
        current_section = "xhs_books"
        continue
    elif "少儿童书——抖音" in t:
        current_section = "dy_children"
        continue
    elif "图书——抖音" in t:
        current_section = "dy_books"
        continue
    
    if current_section in ("xhs_children", "xhs_books"):
        # Xiaohongshu format: "desc - account | 小红书 ..." + URL
        link_match = re.search(r'(https?://[^\s]+)', t)
        url = link_match.group(1) if link_match else ""
        
        # Extract account name: " - XXX | 小红书" or just desc
        account_match = re.search(r'-\s*(.+?)\s*\|', t)
        desc_match = re.search(r'^\d+\s*【(.+?)】', t)
        
        account = account_match.group(1).strip() if account_match else ""
        desc = desc_match.group(1).strip() if desc_match else t[:100]
        
        if url:
            target = xhs_children if current_section == "xhs_children" else xhs_books
            target.append({"account": account, "desc": desc, "url": url})
    
    elif current_section in ("dy_children", "dy_books"):
        if t.startswith("http"):
            # URL line - next paragraph should be description
            current_url = t
            # Check if this line also has description
        else:
            # Description line
            target = dy_children if current_section == "dy_children" else dy_books
            # Find preceding URL from previous paragraph
            pass

# Better parsing for Douyin - pair URLs with descriptions
dy_lines = []
collecting_dy = False
current_dy_section = None
for p in links_doc.paragraphs:
    t = p.text.strip()
    if not t:
        if collecting_dy and dy_lines:
            # Flush
            for i in range(0, len(dy_lines), 2):
                url = dy_lines[i]
                desc = dy_lines[i+1] if i+1 < len(dy_lines) else ""
                if current_dy_section == "dy_children":
                    dy_children.append({"desc": desc, "url": url})
                else:
                    dy_books.append({"desc": desc, "url": url})
            dy_lines = []
            collecting_dy = False
        continue
    
    if "少儿童书——抖音" in t:
        current_dy_section = "dy_children"
        collecting_dy = True
        dy_lines = []
        continue
    elif "图书——抖音" in t:
        current_dy_section = "dy_books"
        collecting_dy = True
        dy_lines = []
        continue
    elif "少儿童书——小红书" in t or "图书——小红书" in t:
        if collecting_dy and dy_lines:
            for i in range(0, len(dy_lines), 2):
                url = dy_lines[i]
                desc = dy_lines[i+1] if i+1 < len(dy_lines) else ""
                if current_dy_section == "dy_children":
                    dy_children.append({"desc": desc, "url": url})
                else:
                    dy_books.append({"desc": desc, "url": url})
            dy_lines = []
        collecting_dy = False
        continue
    
    if collecting_dy:
        dy_lines.append(t)

# Flush remaining
if collecting_dy and dy_lines:
    for i in range(0, len(dy_lines), 2):
        url = dy_lines[i]
        desc = dy_lines[i+1] if i+1 < len(dy_lines) else ""
        if current_dy_section == "dy_children":
            dy_children.append({"desc": desc, "url": url})
        else:
            dy_books.append({"desc": desc, "url": url})

print(f"XHS Children: {len(xhs_children)}")
for x in xhs_children:
    print(f"  {x['account']}: {x['url'][:80]}")
print(f"\nXHS Books: {len(xhs_books)}")
for x in xhs_books:
    print(f"  {x['account']}: {x['url'][:80]}")
print(f"\nDY Children: {len(dy_children)}")
for x in dy_children:
    print(f"  {x['url'][:60]} | {x['desc'][:60]}")
print(f"\nDY Books: {len(dy_books)}")
for x in dy_books:
    print(f"  {x['url'][:60]} | {x['desc'][:60]}")
